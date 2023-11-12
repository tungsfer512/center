from elasticsearch import Elasticsearch
import redis
from rest_framework.pagination import PageNumberPagination
from io import BytesIO
import xlsxwriter
from datetime import datetime, timezone, timedelta
from django.shortcuts import render
from django.http import HttpResponse, FileResponse
from rest_framework.response import Response
from rest_framework import viewsets
from rest_framework.views import APIView
from drf_yasg.utils import swagger_auto_schema
from drf_yasg import openapi
from rest_framework.decorators import api_view
from django_filters.rest_framework import DjangoFilterBackend
from devices.models import *
from rest_framework import status, parsers
from rest_framework.decorators import action
from devices.serializers import *
from rest_framework import permissions
import json
import os
from os.path import isfile
import requests
import logging
import re
from django.core.paginator import Paginator
import time
from django.utils import timezone
from pathlib import Path
from .tools import mail, sms, tele
from CenterServer.env_dev import *
from docx import Document
import subprocess

# Create your views here.

pathEnv = "/CenterServer/.env.dev"


class ModelMachineLearningView(viewsets.ModelViewSet):
    serializer_class = ModelMachineLeaningSerializer
    parser_classes = (parsers.FormParser,
                      parsers.MultiPartParser, parsers.FileUploadParser)
    queryset = ModelML.objects.all()
    permission_classes = [permissions.IsAdminUser]

    # Get latest Model machine learning

    @swagger_auto_schema(
        method='get'
    )
    @action(detail=False, methods=['get'])
    def get_latest_model_ML(self, request):
        try:
            model_ML = ModelML.objects.latest('timestamp')
            file_handle = model_ML.file
            response = HttpResponse(
                file_handle, content_type='application/octet-stream')
            response['Content-Disposition'] = 'attachment; filename="model.h5"'
            return response
        except Exception as e:
            logging.error(e)
            print(e)
            return Response(data={'status': False}, status=status.HTTP_400_BAD_REQUEST)


class IoTAnalyzerDevicesView(viewsets.ModelViewSet):

    serializer_class = DevicesSerializer
    queryset = Devices.objects.all()
    permission_classes = [permissions.IsAdminUser]
    parser_classes = (parsers.FormParser,
                      parsers.MultiPartParser, parsers.FileUploadParser)
    id = openapi.Parameter(
        'id',
        openapi.IN_QUERY,
        description='id',
        type=openapi.TYPE_INTEGER
    )
    auth_token = openapi.Parameter(
        'auth_token',
        openapi.IN_QUERY,
        description='auth_token',
        type=openapi.TYPE_STRING
    )
    connect = openapi.Parameter(
        'connect',
        openapi.IN_QUERY,
        description='connect',
        type=openapi.TYPE_BOOLEAN
    )
    analyzer_ip = openapi.Parameter(
        'analyzer_ip',
        openapi.IN_QUERY,
        description='analyzer_ip',
        type=openapi.TYPE_STRING,
    )
    # Get performance IoTAnalyzer

    @action(detail=False, methods=['get'])
    def get_best_analyzer(self, request):
        try:
            arr_json = []
            redis_center = redis.Redis(
                host=os.getenv('REDIS_SERVER'), port=6379, db=0)
            analyzer_list = redis_center.get("analyzer_list")
            if analyzer_list:
                string_analyzer_list = redis_center.get(
                    "analyzer_list").decode("utf-8").replace("'", '"')
                analyzer_list = json.loads(string_analyzer_list)
            else:
                analyzer_list = []
            for analyzer in analyzer_list:
                cpu_ram = json.loads(
                    redis_center.get(analyzer).decode("utf-8"))
                cpu = cpu_ram["cpu"]
                ram = cpu_ram["ram"]
                latency = cpu_ram["latency"]
                check_receive = cpu_ram["check_receive"]
                data = {"ip": analyzer, "cpu": cpu,
                        "ram": ram, "latency": latency, "check_receive": check_receive}
                if len(Devices.objects.filter(ip=analyzer)) and str(check_receive).lower() == "true":
                # if len(Devices.objects.filter(ip=analyzer)):
                    arr_json.append(data)
            return Response(data=arr_json, status=status.HTTP_200_OK)
        except Exception as e:
            print(e)
            return Response(data={'status': False}, status=status.HTTP_400_BAD_REQUEST)

    # @swagger_auto_schema(
    #     method='post',
    #     request_body=DeviceDistributedThresholdDataSerializer
    # )
    # @action(detail=False, methods=['post'])
    # def post_threshold_analyzer(self, request):
    #     try:
    #         request.data._mutable = True
    #         data = request.data
    #         analyzer_ip = data.get("analyzer_ip", None)
    #         del data['analyzer_ip']
    #         device = Devices.objects.filter(ip=analyzer_ip).first()
    #         data['device'] = device.id
    #         serializer = DeviceDistributedThresholdSerializer(data=data)
    #         if serializer.is_valid():
    #             serializer.save()
    #             return Response(data=serializer.data, status=status.HTTP_201_CREATED)
    #         else:
    #             return Response(data=serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    #     except Exception as e:
    #         return Response(str(e), status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    @swagger_auto_schema(
        method='put',
        request_body=DeviceDistributedThresholdDataSerializer
    )
    @action(detail=False, methods=['put'])
    def put_threshold_analyzer(self, request, *args, **kwargs):
        try:
            data = request.data
            analyzer_ip = data.get("analyzer_ip", None)
            analyzer_update = requests.put(f"http://{analyzer_ip}:8000/devices/threshold", json={
                'latency_threshold': data.get("min_latency", 10),
                'ram_threshold': data.get("min_used_ram", 90),
                'cpu_threshold': data.get("min_load_cpu", 90),
                'active_distributed': data.get("check", True),
                'active_distributed_receive': data.get("check_receive", True),
            })
            if analyzer_update.status_code == 200:
                return Response(data={
                    'latency_threshold': data.get("min_latency", 10),
                    'ram_threshold': data.get("min_used_ram", 90),
                    'cpu_threshold': data.get("min_load_cpu", 90),
                    'active_distributed': data.get("check", True),
                    'active_distributed_receive': data.get("check_receive", True),
                }, status=status.HTTP_200_OK)
        except Exception as e:
            return Response(data=str(e), status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    # @swagger_auto_schema(
    #     method='delete',
    #     manual_parameters=[analyzer_ip]
    # )
    # @action(detail=False, methods=['delete'])
    # def delete_threshold_analyzer(self, request, device = None):
    #     try:
    #         analyzer_ip = request.query_params.dict().get('analyzer_ip', None)
    #         device = Devices.objects.filter(ip=analyzer_ip).first()
    #         DeviceDistributedThreshold.objects.filter(device_id=device.id).delete()

    #         return Response(data="deleted", status=status.HTTP_200_OK)
    #     except:
    #         return Response(data={'status': False}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    @swagger_auto_schema(
        method='get',
        manual_parameters=[analyzer_ip],
    )
    @action(detail=False, methods=['get'])
    def get_threshold_analyzer(self, request):
        try:
            analyzer_ip = request.query_params.dict().get('analyzer_ip', None)
            print("analyzer_ip:", analyzer_ip)
            if analyzer_ip != None:
                resData = []
                device = Devices.objects.all().filter(ip=analyzer_ip).first()
                analyzer_get = requests.get(
                    f"http://{analyzer_ip}:8000/devices/threshold")
                if analyzer_get.status_code == 200:
                    res_te = json.loads(analyzer_get.text)
                    res = {
                        "min_latency": res_te["latency_threshold"],
                        "min_used_ram": res_te["ram_threshold"],
                        "min_load_cpu": res_te["cpu_threshold"],
                        "check": res_te["active_distributed"],
                        "check_receive": res_te["active_distributed_receive"],
                    }
                    res["device"] = DevicesSerializer(device).data
                    resData.append(res)
                    return Response(data=resData, status=status.HTTP_200_OK)
                else:
                    return Response(data={'status': False}, status=status.HTTP_400_BAD_REQUEST)
            else:
                resData = []
                for device in Devices.objects.all():
                    analyzer_get = requests.get(
                        f"http://{device.ip}:8000/devices/threshold")
                    if analyzer_get.status_code == 200:
                        res_te = json.loads(analyzer_get.text)
                        res = {
                            "min_latency": res_te["latency_threshold"],
                            "min_used_ram": res_te["ram_threshold"],
                            "min_load_cpu": res_te["cpu_threshold"],
                            "check": res_te["active_distributed"],
                            "check_receive": res_te["active_distributed_receive"],
                        }
                        res["device"] = DevicesSerializer(device).data
                        resData.append(res)
                    else:
                        return Response(data={'status': False}, status=status.HTTP_400_BAD_REQUEST)
                return Response(data=resData, status=status.HTTP_200_OK)
        except Exception as e:
            return Response(data=str(e), status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            

    @action(detail=False, methods=['get'])
    def get_device_performance(self, request):
        try:
            device = Devices.objects.filter(
                id=request.query_params.get('id')).first()
            ip = device.ip
            port = device.port
            data = {
                "username": device.username,
                "password": device.password
            }
            api_url_token = f'http://{ip}:{port}/auth/token/login'

            req = requests.post(api_url_token, json=data)
            json_data = req.json()
            token = json_data['auth_token']
            header = f'Token {token}'
            api_url = f'http://{ip}:{port}/dashboard/stat/'
            device_data = requests.get(
                api_url, headers={'Authorization': header}).json()
            return Response(data=device_data, status=status.HTTP_201_CREATED)
        except Exception as e:
            logging.error(e)
            print(e)
            return Response(data={'status': False}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method='post',
        manual_parameters=[id]
    )
    @action(detail=False, methods=['post'])
    def add_analyzer_center_server(self, request):
        try:
            device = Devices.objects.filter(
                id=request.query_params.get('id')).first()
            ip = device.ip
            port = device.port
            data = {
                "username": device.username,
                "password": device.password
            }
            api_url_token = f'http://{ip}:{port}/auth/token/login'

            req = requests.post(api_url_token, json=data)
            json_data = req.json()
            token = json_data['auth_token']
            header = f'Token {token}'
            api_url = f'http://{ip}:{port}/CenterServer/'
            create_new_center = requests.post(api_url, data={'ip': os.getenv(
                'django_ip'), 'connect': False, 'auth': ''}, headers={'Authorization': header})
        except Exception as e:
            logging.error(e)
            print(e)
            return Response(data={'status': False}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method='put',
        manual_parameters=[id, auth_token, connect]
    )
    @action(detail=False, methods=['put'])
    def update_alert_analyzer(self, request):
        try:
            device = Devices.objects.filter(
                id=request.query_params.get('id')).first()
            ip = device.ip
            port = device.port
            data = {
                "username": device.username,
                "password": device.password
            }
            api_url_token = f'http://{ip}:{port}/auth/token/login'

            req = requests.post(api_url_token, json=data)
            json_data = req.json()
            token = json_data['auth_token']
            header = f'Token {token}'
            api_url = f'http://{ip}:{port}/CenterServer/update_by_ip/'
            update_center = requests.put(api_url, data={'ip': os.getenv('django_ip'), 'connect': request.query_params.get(
                'connect'), 'auth': request.query_params.get('auth_token')}, headers={'Authorization': header})
        except Exception as e:
            logging.error(e)
            print(e)
            return Response(data={'status': False}, status=status.HTTP_400_BAD_REQUEST)


class BlackListIPView(viewsets.ModelViewSet):

    serializer_class = BlackListIPSerializer
    queryset = BlackListIP.objects.all()
    permission_classes = [permissions.IsAdminUser]

    ip = openapi.Parameter(
        'ip',
        openapi.IN_QUERY,
        description='IP',
        type=openapi.TYPE_STRING,
    )
    page_size = openapi.Parameter(
        'page_size',
        openapi.IN_QUERY,
        description='Số lượng mỗi trang',
        type=openapi.TYPE_STRING,
        required=True
    )
    page = openapi.Parameter(
        'page',
        openapi.IN_QUERY,
        description='Trang',
        type=openapi.TYPE_STRING,
        required=True
    )
    url = openapi.Parameter(
        'url',
        openapi.IN_QUERY,
        description='Địa chỉ trang web',
        type=openapi.TYPE_STRING,
    )

    @swagger_auto_schema(
        method='get',
        manual_parameters=[ip, url, page, page_size],
    )
    @action(detail=False, methods=['get'])
    def find_black_list_IP_in_list(self, request):
        try:
            data = BlackListIP.objects.all()
            if request.query_params.get('ip', None) != None:
                data = data.filter(
                    ip__icontains=request.query_params.get('ip', None)
                )
            if request.query_params.get('url', None) != None:
                data = data.filter(
                    url__icontains=request.query_params.get('url', None)
                )
            page_size = request.query_params['page_size']
            page = request.query_params['page']
            paginator = Paginator(data.order_by('id'), page_size)
            serializer = self.get_serializer(
                paginator.get_page(page), many=True)
            return Response({
                'results': serializer.data,
                'count': paginator.count,
            })
        except Exception as e:
            logging.error(e)
            return Response(data={'status': str(e)}, status=status.HTTP_400_BAD_REQUEST)


class WhiteListIPView(viewsets.ModelViewSet):

    serializer_class = WhiteListIPSerializer
    queryset = WhiteListIP.objects.all()
    permission_classes = [permissions.IsAdminUser]

    ip = openapi.Parameter(
        'ip',
        openapi.IN_QUERY,
        description='IP',
        type=openapi.TYPE_STRING,
    )
    page_size = openapi.Parameter(
        'page_size',
        openapi.IN_QUERY,
        description='Số lượng mỗi trang',
        type=openapi.TYPE_STRING,
        required=True
    )
    page = openapi.Parameter(
        'page',
        openapi.IN_QUERY,
        description='Trang',
        type=openapi.TYPE_STRING,
        required=True
    )
    url = openapi.Parameter(
        'url',
        openapi.IN_QUERY,
        description='Địa chỉ trang web',
        type=openapi.TYPE_STRING,
    )

    @swagger_auto_schema(
        method='get',
        manual_parameters=[ip, url, page, page_size],
    )
    @action(detail=False, methods=['get'])
    def find_white_list_IP_in_list(self, request):
        try:
            data = WhiteListIP.objects.all()
            if request.query_params.get('ip', None) != None:
                data = data.filter(
                    ip__icontains=request.query_params.get('ip', None)
                )
            if request.query_params.get('url', None) != None:
                data = data.filter(
                    url__icontains=request.query_params.get('url', None)
                )
            page_size = request.query_params['page_size']
            page = request.query_params['page']
            paginator = Paginator(data.order_by('id'), page_size)
            serializer = self.get_serializer(
                paginator.get_page(page), many=True)
            return Response({
                'results': serializer.data,
                'count': paginator.count,
            })
        except Exception as e:
            logging.error(e)
            return Response(data={'status': str(e)}, status=status.HTTP_400_BAD_REQUEST)


class StandardResultsSetPagination(PageNumberPagination):
    page_size = 10
    page_size_query_param = 'page_size'
    max_page_size = 1000


class AlertsViewSet(viewsets.ModelViewSet):
    queryset = Alerts.objects.all().order_by('-timestamp', 'status')
    # queryset = get_devices.delay().get()
    serializer_class = AlertsSerializer
    # permission_classes = [permissions.IsAuthenticated]
    # logger.error('================================', get_devices.delay())
    pagination_class = StandardResultsSetPagination

    filter_backends = [DjangoFilterBackend]
    filterset_fields = ['ip', 'timestamp',
                        'hash', 'pid', 'message', 'type', 'status', 'device']

    def create(self, request, *args, **kwargs):
        try:
            print("override create alert !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
            
            data = request.data
            print(data)
            loai = data.get("loai", "")
            local = data.get("local", "")
            device_ip = data.get("device_ip")
            device = Devices.objects.filter(ip=device_ip).first()
            type_log = data.get("type_log", "alert_attack_kc")
            
            del data["loai"]
            del data["local"]
            del data["device_ip"]
            del data["type_log"]
            data['device'] = device.id
            
            es = Elasticsearch(hosts=f"{os.environ.get('ELASTIC_SERVER', 'http://es-container-center:9200')}")
            doc = {
                "description": f"{data.get('message')}",
                "timestamp": datetime.datetime.now(timezone.utc),
                "loai": loai,
                "device_id": data.get("device"),
                'ip_attack': data.get("ip"),
                'local': local,
                "type_log": type_log,
            }
            es.index(index="demo-kc", document=doc)
            print("override create alert send elastic !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
            
            serializer = AlertsSerializer(data = data)
            if serializer.is_valid():
                serializer.save()
            else:
                print("-----___________serializer.errors", serializer.errors)
                return Response(serializer.errors, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
            return Response(serializer.data, status=status.HTTP_201_CREATED)
            
        except Exception as e:
            return Response(str(e), status=status.HTTP_500_INTERNAL_SERVER_ERROR)    

    ip = openapi.Parameter(
        'ip',
        openapi.IN_QUERY,
        description='IP thiết bị',
        type=openapi.TYPE_STRING,

    )
    device = openapi.Parameter(
        'device',
        openapi.IN_QUERY,
        description='ID thiết bị',
        type=openapi.TYPE_STRING,

    )
    page = openapi.Parameter(
        'current',
        openapi.IN_QUERY,
        description='Trang',
        type=openapi.TYPE_STRING,
        required=True
    )
    page_size = openapi.Parameter(
        'page_size',
        openapi.IN_QUERY,
        description='Số lượng mỗi trang',
        type=openapi.TYPE_STRING,
        required=True
    )
    Type = openapi.Parameter(
        'type',
        openapi.IN_QUERY,
        description='Loại thông báo',
        type=openapi.TYPE_STRING,
        # required = True

    )
    status = openapi.Parameter(
        'status',
        openapi.IN_QUERY,
        description='Trạng thái',
        type=openapi.TYPE_STRING,
        # required = True
    )

    @swagger_auto_schema(
        method='get',
        manual_parameters=[ip, page, page_size, Type, status],
        # responses=openapi.Response('AlertsSerializer', AlertsSerializer),
    )
    @action(detail=False, methods=['get'])
    def GetAlertsByIP(self, request):
        try:
            type = request.query_params.get('type', None)
            statusValue = request.query_params.get('status', None)
            filters = {"ip": request.query_params['ip']}
            if type:
                filters["type"] = type
            if statusValue:
                filters["status"] = statusValue
            result = Alerts.objects.filter(**filters).order_by('-timestamp')

            page_size = request.query_params['page_size']
            page = request.query_params['current']

            paginator = Paginator(result, page_size)
            serializer = self.get_serializer(
                paginator.get_page(page), many=True)
            return Response({
                'results': serializer.data,
                'count': paginator.count,
            })
        except Exception as e:
            logging.error(e)
            return Response(data={'status': False}, status=status.HTTP_400_BAD_REQUEST)

    device_id = openapi.Parameter(
        'device_id',
        openapi.IN_QUERY,
        description='device_id',
        required=True,
        type=openapi.TYPE_STRING,
        # required = True
    )
    pid = openapi.Parameter(
        'pid',
        openapi.IN_QUERY,
        description='pid',
        required=True,
        type=openapi.TYPE_STRING,
        # required = True
    )
    alert_id = openapi.Parameter(
        'alert_id',
        openapi.IN_QUERY,
        description='alert_id',
        required=True,
        type=openapi.TYPE_STRING,
        # required = True
    )
    message = openapi.Parameter(
        'message',
        openapi.IN_QUERY,
        description='message',
        type=openapi.TYPE_STRING,
        # required = True
    )
    all_device = openapi.Parameter(
        'all_device',
        openapi.IN_QUERY,
        description='all_device',
        type=openapi.TYPE_BOOLEAN,
        required=True
    )

    @swagger_auto_schema(
        method='patch',
        manual_parameters=[message, ip, page_size, page, all_device],
        request_body=AlertsExportSerializer
    )
    @action(detail=False, methods=['patch'])
    def search_alert(self, request):
        try:
            print("check 1")
            xxx_data = json.loads(request.query_params['data'])
            print("check 2")
            result = Alerts.objects.all()
            print("check 3")
            if request.query_params.get('ip', None) != None:
                result = result.filter(
                    ip__icontains=request.query_params['ip']
                )
            print("check 4")
            if request.query_params.get('message', None) != None:
                result = result.filter(
                    message__icontains=request.query_params['message']
                )
            if request.query_params.get('address', None) != None:
                result = result.filter(
                    address__icontains=request.query_params['address']
                )
            print("check 5")
            print("check 5")

            if xxx_data.get('timestamp', None) != None:
                print(type(result))
                print(xxx_data.get('timestamp', None))
                # aaa = datetime.strptime(xxx_data.get('timestamp', None)[0], '%yyyy-%mm-%dd %HH:%MM:%SS.')
                # bbb = datetime.strptime(xxx_data.get('timestamp', None)[1], '%yyyy-%mm-%dd %HH:%MM:%SS.')
                # print(aaa, bbb)
                result = result.filter(timestamp__range=[xxx_data.get(
                    'timestamp', None)[0], xxx_data.get('timestamp', None)[1]])
            i = 0
            xyz = result
            if request.query_params['all_device'] == 'false':
                if xxx_data.get('device', None) != None:
                    for id in xxx_data.get('device', None):
                        if i == 0:
                            xyz = result.filter(device=id)
                            abc = xyz
                            i = 1

                        else:
                            # xyz.union(abc, result.filter(device = id))
                            xyz = (abc | result.filter(device=id)).distinct()
                            abc = xyz

            result = xyz
            i = 0

            if xxx_data.get('type', None) != None:
                for id in xxx_data.get('type', None):
                    if i == 0:
                        xyz = result.filter(type=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = (abc | result.filter(type=id))
                        abc = xyz

            result = xyz
            i = 0
            if xxx_data.get('status', None) != None:
                for id in xxx_data.get('status', None):
                    if i == 0:
                        xyz = result.filter(status=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = (abc | result.filter(status=id)).distinct()
                        abc = xyz

            print("check 6")

            page_size = request.query_params['page_size']
            page = request.query_params['current']
            paginator = Paginator(xyz, page_size)
            # print(paginator.get_page(page))
            serializer = self.get_serializer(
                paginator.get_page(page), many=True)
            # print(serializer)
            return Response({
                'results': serializer.data,
                'count': paginator.count,
            })

        except Exception as e:
            logging.error(e)
            return Response(data={'status': False, 'Log bug': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        method='patch',
        request_body=AlertsExportSerializer,
        manual_parameters=[message, ip, all_device],
        # manual_parameters=[message],
    )
    @action(detail=False, methods=['patch'])
    def exportXLSX(self, request):
        try:
            xxx_data = json.loads(request.query_params['data'])
            print(xxx_data)
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            sheet1 = workbook.add_worksheet("Sheet1")
            # sheet1 = workbook.add_sheet('Sheet1', cell_overwrite_ok = True)
            col_stt = 1
            rol = ["A", "B", "C", "D", "E", "F", "G",
                   "H", "I", "J", "K", "L", "M", "N"]
            rol_stt = 0
            for i in xxx_data["status"]:
                sheet1.write(
                    # rol[rol_stt]
                    "A" + str(col_stt), "Số lần thực hiện của thể loại : " + i)
                # col_stt = col_stt + 1
                sheet1.write(
                    # rol[rol_stt]
                    "B" + str(col_stt), "Số lần thực hiện là " + str(Alerts.objects.filter(
                        status=i
                    ).count())
                )
                col_stt = col_stt + 1
            for y in xxx_data["type"]:
                sheet1.write(
                    # rol[rol_stt]
                    "A" + str(col_stt), "Số lần thực hiện của trạng thái: " + y)
                # col_stt = col_stt + 1
                sheet1.write(rol[rol_stt+1] + str(col_stt), "Số lần thực hiện là " + str(Alerts.objects.filter(
                    type=y
                ).count())
                )
                col_stt = col_stt + 1
            sheet2 = workbook.add_worksheet("Sheet2")
            result = Alerts.objects.all()

            if request.query_params.get('ip', None) != None:
                result = result.filter(
                    ip__icontains=request.query_params['ip']
                )

            if request.query_params.get('message', None) != None:
                result = result.filter(
                    message__icontains=request.query_params['message']
                )
            if request.query_params.get('address', None) != None:
                result = result.filter(
                    address__icontains=request.query_params['address']
                )

            if xxx_data.get('timestamp', None) != None:
                result = result.filter(timestamp__gte=xxx_data.get('timestamp', None)[0],
                                        timestamp__lte=xxx_data.get('timestamp', None)[1])
            i = 0
            result_id = result
            if request.query_params['all_device'] == True:
                if xxx_data.get('id', None) != None:
                    for id in xxx_data.get('id', None):
                        if i == 0:
                            xyz = result.filter(device=id)
                            abc = xyz
                            i = 1

                        else:
                            # xyz.union(abc, result.filter(device = id))
                            xyz = (abc | result.filter(device=id)).distinct()
                            abc = xyz

            result = xyz
            i = 0
            if xxx_data.get('type', None) != None:
                for id in xxx_data.get('type', None):
                    if i == 0:
                        xyz = result.filter(type=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = (abc | result.filter(type=id)).distinct()
                        abc = xyz

            result = xyz
            i = 0
            if xxx_data.get('status', None) != None:
                for id in xxx_data.get('status', None):
                    if i == 0:
                        xyz = result.filter(status=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = (abc | result.filter(status=id)).distinct()
                        abc = xyz

            header = ["STT", "ID", "IP", "Timestamp", "Hash", "PID",
                      "Message", "Type", "Rule_base", "Log", "Status", "Device"]
            col_stt = 1
            for rol_stt in range(12):
                sheet2.write(
                    rol[rol_stt] + str(col_stt), header[rol_stt])
            col_stt = col_stt + 1
            for dt_res in result_id:

                sheet2.write(rol[0] + str(col_stt), col_stt - 1)
                sheet2.write(rol[1] + str(col_stt), dt_res.id)
                sheet2.write(rol[2] + str(col_stt), dt_res.ip)
                sheet2.write(rol[3] + str(col_stt),
                             dt_res.timestamp.strftime('%Y-%m-%d %H:%M:%S'))
                sheet2.write(rol[4] + str(col_stt), dt_res.hash)
                sheet2.write(rol[5] + str(col_stt), dt_res.pid)
                sheet2.write(rol[6] + str(col_stt), dt_res.message)
                sheet2.write(rol[7] + str(col_stt), dt_res.type)
                sheet2.write(rol[8] + str(col_stt), dt_res.rule_base)
                sheet2.write(rol[9] + str(col_stt), dt_res.log)
                sheet2.write(rol[10] + str(col_stt), dt_res.status)
                sheet2.write(rol[11] + str(col_stt), dt_res.device.id)
                col_stt = col_stt + 1
            workbook.close()

            # create a response
            response = HttpResponse(content_type='application/vnd.ms-excel')

            # tell the browser what the file is named
            response['Content-Disposition'] = 'attachment;filename="some_file_name.xlsx"'

            # put the spreadsheet data into the response
            response.write(output.getvalue())

            return response

        except Exception as e:
            logging.error(e)
            return Response(data={'status': False, 'Log bug': e}, status=status.HTTP_400_BAD_REQUEST)


@api_view(('GET',))
def _retrieve_or_create_mail(request, *args, **kwargs):
    queryset = DevicesMailCheck.objects.all()
    path_variable = kwargs
    if request.method == 'GET':
        try:
            device = path_variable.get('device_id')
            print(device)
            mailcheck = queryset.filter(device=device).first()
            if Devices.objects.all().filter(id=device).first() == None:
                return Response({
                    'message': "Device not found!"
                }, status=status.HTTP_404_NOT_FOUND)
            if mailcheck != None:
                print("check not null")
                return Response({
                    'data': DevicesMailCheckSerializer(mailcheck).data
                }, status=status.HTTP_200_OK)
            else:
                print("check null")
                mailcheckserializer = DevicesMailCheckSerializer(data={
                    'network': True,
                    'malware': True,
                    'syscall': True,
                    'log': True,
                    'ips': True,
                    'device': device
                })
                if mailcheckserializer.is_valid():
                    print("check null valid")
                    # print(mailcheckserializer.data)
                    mailcheckserializer.save()
                else:
                    print(mailcheckserializer.errors)
                return Response({
                    'data': mailcheckserializer.data,
                }, status=status.HTTP_200_OK)
        except Exception as error:
            return Response({
                'message': str(error)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@swagger_auto_schema(
    method='put',
    request_body=DevicesMailCheckSerializer
)
@api_view(('PUT',))
def _update_mail(request, *args, **kwargs):
    queryset = DevicesMailCheck.objects.all()
    data = request.data
    if request.method == 'PUT':
        try:

            mailcheck = queryset.filter(device=data.get('device')).first()
            print(mailcheck)
            mailcheckserializer = DevicesMailCheckSerializer(
                mailcheck, data=data)
            if mailcheckserializer.is_valid():
                mailcheckserializer.save()
            return Response({
                'data': mailcheckserializer.data,
            }, status=status.HTTP_200_OK)
        except Exception as error:
            return Response({
                'message': str(error)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(("POST", ))
def _send_mail(request, *args, **kwargs):
    try:
        data = request.data
        if data.get("device_ip", None) == None:
            time_auto = int(get_env("AUTO_SEND_MAIL_SECOND", 60 * 60 * 24))
            if time_auto == -1:
                time_auto = 60 * 60 * 24
            # print("time111111111111111111111111111111111, ", time_auto)
            days = time_auto // (60 * 60 * 24)
            time_auto -= (60 * 60 * 24 * days)
            hours = time_auto // (60 * 60)
            time_auto -= (60 * 60 * hours)
            mins = time_auto // (60)
            time_auto -= (60 * mins)
            secs = time_auto
            analyzers = Devices.objects.all()
            time_start = datetime.datetime.strftime(datetime.datetime.now() - timedelta(days=days, hours=hours, minutes=mins, seconds=secs), '%Y-%m-%d %H:%M:%S.%fZ')
            time_end = datetime.datetime.strftime(datetime.datetime.now(), '%Y-%m-%d %H:%M:%S.%fZ')
            print("===============================================")
            print(time_start, time_end)
            print("===============================================")
            for analyzer in analyzers:
                # result = Alerts.objects.all()
                result = Alerts.objects.all().filter(timestamp__range=(time_start, time_end))
                            
                output = "./devices/attach.xlsx"
                workbook = xlsxwriter.Workbook(output)
                worksheet = workbook.add_worksheet()
                worksheet.write('A1', 'STT')
                worksheet.write('B1', 'Mã thiết bị')
                worksheet.write('C1', 'Địa chỉ ip')
                worksheet.write('D1', 'Cảnh báo')
                worksheet.write('E1', 'Hash')
                worksheet.write('F1', 'PID')
                worksheet.write('G1', 'Thời gian')
                for i in range(len(result)):
                    worksheet.write(f'A{i + 2}', i + 1)
                    worksheet.write(f'B{i + 2}', result[i].device_id)
                    worksheet.write(f'C{i + 2}', result[i].ip)
                    worksheet.write(f'D{i + 2}', result[i].message)
                    worksheet.write(f'E{i + 2}', result[i].hash)
                    worksheet.write(f'F{i + 2}', result[i].pid)
                    worksheet.write(f'G{i + 2}', str(result[i].timestamp))
                workbook.close()
                    
                # subject = f'''Thông tin tổng hợp các cuộc tấn công vào mạng lưới thiết bị IoT từ {time_start} đến {time_end}'''
                subject = f'''Thông tin tổng hợp các cuộc tấn công vào mạng lưới thiết bị IoT từ '''
                body = f'''Trong vòng {days} ngày {hours} giờ {mins} phút {secs} giây vừa qua, đã có tổng cộng {len(result)} cuộc tấn công vào mạng lưới thiết bị IoT của bạn. Vui lòng kiểm tra tệp đính kèm để biết thêm chi tiết.'''
                recipients = [analyzer.email]
                # if len(result) > 0:
                print("+++++++++++++++++++++++++++++++++++++++++++++")
                mail.send_email(subject=subject, body=body, recipients=recipients)
        else:
            device_ip = data.get("device_ip")
            device = Devices.objects.filter(ip=device_ip).first()
            subject = data.get("subject")
            body = data.get("body")
            recipients = [ device.email ]
            mail.send_email(subject=subject, body=body, recipients=recipients)
        return Response({
            'message': "Sent email"
        }, status=status.HTTP_200_OK)
    except Exception as e:
        logging.error(e)
        return Response(data={'status': False, 'Log bug': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(('GET',))
def _retrieve_or_create_sms(request, *args, **kwargs):
    queryset = DevicesSMSCheck.objects.all()
    path_variable = kwargs
    if request.method == 'GET':
        try:
            device = path_variable.get('device_id')
            print(device)
            smscheck = queryset.filter(device=device).first()
            if Devices.objects.all().filter(id=device).first() == None:
                return Response({
                    'message': "Device not found!"
                }, status=status.HTTP_404_NOT_FOUND)
            if smscheck != None:
                print("check not null")
                return Response({
                    'data': DevicesSMSCheckSerializer(smscheck).data
                }, status=status.HTTP_200_OK)
            else:
                print("check null")
                smscheckserializer = DevicesSMSCheckSerializer(data={
                    'network': True,
                    'malware': True,
                    'syscall': True,
                    'log': True,
                    'ips': True,
                    'device': device
                })
                if smscheckserializer.is_valid():
                    print("check null valid")
                    # print(smscheckserializer.data)
                    smscheckserializer.save()
                else:
                    print(smscheckserializer.errors)
                return Response({
                    'data': smscheckserializer.data,
                }, status=status.HTTP_200_OK)
        except Exception as error:
            return Response({
                'message': str(error)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@swagger_auto_schema(
    method='put',
    request_body=DevicesSMSCheckSerializer
)
@api_view(('PUT',))
def _update_sms(request, *args, **kwargs):
    queryset = DevicesSMSCheck.objects.all()
    data = request.data
    if request.method == 'PUT':
        try:

            smscheck = queryset.filter(device=data.get('device')).first()
            print(smscheck)
            smscheckserializer = DevicesSMSCheckSerializer(smscheck, data=data)
            if smscheckserializer.is_valid():
                smscheckserializer.save()
            return Response({
                'data': smscheckserializer.data,
            }, status=status.HTTP_200_OK)
        except Exception as error:
            return Response({
                'message': str(error)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(("POST", ))
def _send_sms(request, *args, **kwargs):
    try:
        data = request.data
        if data.get("device_ip", None) == None:
            time_auto = int(get_env("AUTO_SEND_SMS_SECOND", 60 * 60 * 24))
            if time_auto == -1:
                time_auto = 60 * 60 * 24
            # print("time111111111111111111111111111111111, ", time_auto)
            days = time_auto // (60 * 60 * 24)
            time_auto -= (60 * 60 * 24 * days)
            hours = time_auto // (60 * 60)
            time_auto -= (60 * 60 * hours)
            mins = time_auto // (60)
            time_auto -= (60 * mins)
            secs = time_auto
            time_start = datetime.datetime.strftime(datetime.datetime.now() - timedelta(days=days, hours=hours, minutes=mins, seconds=secs), '%Y-%m-%d %H:%M:%S.%fZ')
            time_end = datetime.datetime.strftime(datetime.datetime.now(), '%Y-%m-%d %H:%M:%S.%fZ')
            analyzers = Devices.objects.all()
            for analyzer in analyzers:
                result = Alerts.objects.all().filter(timestamp__range=(time_start, time_end))

                body = f'''Trong vòng {days} ngày {hours} giờ {mins} phút {secs} giây vừa qua, đã có tổng cộng {len(result)} cuộc tấn công vào mạng lưới thiết bị IoT của bạn. Vui lòng kiểm tra email để biết thêm chi tiết.'''

                recipients = get_env(
                    'RECEIVER_PHONE_PREFIX') + analyzer.phone[1:]
                sender = get_env("SENDER_PHONE_PREFIX", "+1") + \
                    get_env("SENDER_PHONE", "5076205535")
                # if len(result) > 0:
                print("+++++++++++++++++++++++++++++++++++++++++++++")
                sms.send_sms(from_=sender, to_=recipients, body_=body)
        else:
            device_ip = data.get("device_ip")
            device = Devices.objects.filter(ip=device_ip).first()
            body = data.get("body")
            recipients = get_env(
                'RECEIVER_PHONE_PREFIX') + device.phone[1:]
            sender = get_env("SENDER_PHONE_PREFIX", "+1") + \
                get_env("SENDER_PHONE", "5076205535")
            sms.send_sms(from_=sender, to_=recipients, body_=body)
        return Response({
            'message': f"Sent sms"
        }, status=status.HTTP_200_OK)
    except Exception as error:
        return Response({
            'message': str(error)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


def get_headers(analyzer):
    url = f'{analyzer.domain}/auth/token/login/'
    username = analyzer.username
    password = analyzer.password
    headers = {
        'Content-Type': 'application/json',
        'Accept': 'application/json',
    }
    x = requests.post(
        url, data=json.dumps({'username': username, 'password': password}), headers=headers)
    content = json.loads(x.content)

    auth_token = content['auth_token']
    headers = {
        "Authorization": f"Token {auth_token}",
        'Content-Type': 'application/json',
        'Accept': 'application/json',
    }
    return headers


@api_view(("GET", ))
def _get_device(request, *args, **kwargs):
    try:
        params = request.query_params
        analyzer_id = kwargs.get('analyzer_id')
        analyzer = Devices.objects.all().filter(id=analyzer_id).first()
        headers = get_headers(analyzer=analyzer)
        response = requests.get(
            f"{analyzer.domain}/devices/", headers=headers, json=params)
        return Response(response.json(), status=status.HTTP_200_OK)
    except Exception as error:
        return Response({
            'message': str(error)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(("GET", ))
def _get_dashboard_stat(request, *args, **kwargs):
    try:
        params = request.query_params
        analyzer_id = kwargs.get('analyzer_id')
        analyzer = Devices.objects.all().filter(id=analyzer_id).first()
        headers = get_headers(analyzer=analyzer)
        response = requests.get(
            f"{analyzer.domain}/dashboard/stat/", headers=headers, json=params)
        return Response(response.json(), status=status.HTTP_200_OK)
    except Exception as error:
        return Response({
            'message': str(error)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(("GET", ))
def _get_iptracking_stat_ips(request, *args, **kwargs):
    try:
        params = request.query_params
        analyzer_id = kwargs.get('analyzer_id')
        analyzer = Devices.objects.all().filter(id=analyzer_id).first()
        headers = get_headers(analyzer=analyzer)
        response = requests.get(
            f"{analyzer.domain}/IpsTracking/stat_ips/", headers=headers, json=params)
        return Response(response.json(), status=status.HTTP_200_OK)
    except Exception as error:
        return Response({
            'message': str(error)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class GetExportXLSURL(APIView):
    def get(self, request):
        url = request.get_full_path()
        url = url[:url.find("/url")] + url[url.find("/url") + 4:]
        print("checkkkkkk", url)
        return Response(url)


class ExportXLS(APIView):
    def get(self, request):
        # create our spreadsheet.  I will create it in memory with a StringIO
        try:
            xxx_data = json.loads(request.query_params.get('data', None))
            result = Alerts.objects.all()
            if request.query_params.get('ip', None) != None:
                result = result.filter(
                    ip__icontains=request.query_params['ip']
                )
            if request.query_params.get('message', None) != None:
                result = result.filter(
                    message__icontains=request.query_params['message']
                )
            if request.query_params.get('address', None) != None:
                result = result.filter(
                    address__icontains=request.query_params['address']
                )

            if xxx_data.get('timestamp', None) != None:
                print(type(result))
                print(xxx_data.get('timestamp', None))
                # aaa = datetime.strptime(xxx_data.get('timestamp', None)[0], '%yyyy-%mm-%dd %HH:%MM:%SS.')
                # bbb = datetime.strptime(xxx_data.get('timestamp', None)[1], '%yyyy-%mm-%dd %HH:%MM:%SS.')
                # print(aaa, bbb)
                result = result.filter(timestamp__range=(xxx_data.get(
                    'timestamp', None)[0], xxx_data.get('timestamp', None)[1]))
            i = 0
            xyz = result
            if request.query_params['all_device'] == 'false':
                if xxx_data.get('device', None) != None:
                    for id in xxx_data.get('device', None):
                        if i == 0:
                            xyz = result.filter(device=id)
                            abc = xyz
                            i = 1

                        else:
                            # xyz.union(abc, result.filter(device = id))
                            xyz = (abc | result.filter(device=id)).distinct()
                            abc = xyz

            result = xyz
            i = 0

            if xxx_data.get('type', None) != None:
                for id in xxx_data.get('type', None):
                    if i == 0:
                        xyz = result.filter(type=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = (abc | result.filter(type=id))
                        abc = xyz

            result = xyz
            i = 0
            if xxx_data.get('status', None) != None:
                for id in xxx_data.get('status', None):
                    if i == 0:
                        xyz = result.filter(status=id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = (abc | result.filter(status=id)).distinct()
                        abc = xyz
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            worksheet = workbook.add_worksheet()
            worksheet.write('A1', 'STT')
            worksheet.write('B1', 'Mã thiết bị')
            worksheet.write('C1', 'Địa chỉ ip')
            worksheet.write('D1', 'Cảnh báo')
            worksheet.write('E1', 'Hash')
            worksheet.write('F1', 'PID')
            worksheet.write('G1', 'Thời gian')
            worksheet.write('H1', 'Địa chỉ')
            for i in range(len(xyz)):
                worksheet.write(f'A{i + 2}', i + 1)
                worksheet.write(f'B{i + 2}', xyz[i].device_id)
                worksheet.write(f'C{i + 2}', xyz[i].ip)
                worksheet.write(f'D{i + 2}', xyz[i].message)
                worksheet.write(f'E{i + 2}', xyz[i].hash)
                worksheet.write(f'F{i + 2}', xyz[i].pid)
                worksheet.write(f'G{i + 2}', str(xyz[i].timestamp))
                worksheet.write(f'H{i + 2}', xyz[i].address)
            workbook.close()

            # create a response
            response = HttpResponse(content_type='application/vnd.ms-excel')

            # tell the browser what the file is named
            response['Content-Disposition'] = 'attachment;filename="alerts.xlsx"'

            # put the spreadsheet data into the response
            response.write(output.getvalue())
            return response
        except Exception as e:
            logging.error(e)
            return Response(data={'status': False, 'Log bug': str(e)}, status=status.HTTP_400_BAD_REQUEST)


class GetExportPDFURL(APIView):
    def get (self, request):
        url = request.get_full_path()
        url = url[:url.find("/url")] + url[url.find("/url") + 4:]
        print("checkkkkkk", url)
        return Response(url)

class ExportPDF(APIView):
    def get(self, request):
        
        try:
            xxx_data = json.loads(request.query_params.get('data', None))
            result = Alerts.objects.all()
            if request.query_params.get('ip', None) != None:
                result = result.filter(
                ip__icontains=request.query_params['ip']
                )
            if request.query_params.get('message', None) != None:
                result = result.filter(
                    message__icontains=request.query_params['message']
                )
            if request.query_params.get('address', None) != None:
                result = result.filter(
                    address__icontains=request.query_params['address']
                )
             
            if xxx_data.get('timestamp', None) != None:
                print(type(result))
                print(xxx_data.get('timestamp', None))
                result = result. filter(timestamp__range = (xxx_data.get('timestamp', None)[0], xxx_data.get('timestamp', None)[1]))
            i = 0
            xyz = result
            if request.query_params['all_device'] == 'false':
                if xxx_data.get('device', None) != None:
                    for id in  xxx_data.get('device', None):
                        if i == 0:
                            xyz = result.filter(device = id)
                            abc = xyz
                            i = 1

                        else:
                            # xyz.union(abc, result.filter(device = id))
                            xyz = (abc | result.filter(device = id)).distinct()
                            abc = xyz

            result = xyz
            i=0
            
            if xxx_data.get('type', None) != None:
                for id in xxx_data.get('type', None):
                    if i == 0:
                        xyz = result.filter(type = id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = (abc | result.filter(type = id))
                        abc = xyz
            
            result = xyz
            i=0
            if xxx_data.get('status', None) != None:
                for id in  xxx_data.get('status', None):
                    if i == 0:
                        xyz = result.filter(status = id)
                        abc = xyz
                        i = 1
                    else:
                        # xyz.union(abc, result.filter(type = id))
                        xyz = (abc | result.filter(status = id)).distinct()
                        abc = xyz
            
            buffer = BytesIO()
            document = Document()

            document.add_heading('Cảnh báo tấn công mạng', 0)
            
            table = document.add_table(rows=1, cols=8)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'STT'
            hdr_cells[1].text = 'Mã thiết bị'
            hdr_cells[2].text = 'Địa chỉ ip'
            hdr_cells[3].text = 'Cảnh báo'
            hdr_cells[4].text = 'Hash'
            hdr_cells[5].text = 'PID'
            hdr_cells[6].text = 'Thời gian'
            hdr_cells[7].text = 'Địa chỉ'
            
            for i in range(len(xyz)):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i + 1)
                row_cells[1].text = str(xyz[i].device_id)
                row_cells[2].text = str(xyz[i].ip)
                row_cells[3].text = str(xyz[i].message)
                row_cells[4].text = str(xyz[i].hash)
                row_cells[5].text = str(xyz[i].pid)
                row_cells[6].text = str(xyz[i].timestamp.strftime("%d/%m/%Y, %H:%M:%S"))
                row_cells[7].text = str(xyz[i].address)
            
            table.style = 'Colorful List'
            
            document.save("./devices/reports/alerts.docx")
            convert_to_pdf = f"libreoffice --headless --convert-to pdf ./devices/reports/alerts.docx --outdir ./devices/reports/pdf"
            subprocess.run(convert_to_pdf, shell=True)
            print(os.listdir("./devices/reports/pdf"))
            return FileResponse(open("./devices/reports/pdf/alerts.pdf", 'rb+'), as_attachment=True, filename='alerts.pdf', status=status.HTTP_200_OK)
        except Exception as e:
            logging.error(e)
            return Response(data={'status': False, 'Log bug': str(e)}, status=status.HTTP_400_BAD_REQUEST)


class DistribtutedView(APIView):
    def get(self, request):
        try:
            elastic_server = os.environ.get(
                "ELASTIC_SERVER", "http://192.168.10.162:9200")
            req = requests.get(
                f"{elastic_server}/demo-kc/_search?q=type_log:distribution", )
            datas = req.json().get("hits").get("hits")
            devices = DevicesSerializer(Devices.objects.all(), many=True).data
            distributions = []
            time_now = datetime.datetime.now(timezone.utc).timestamp()
            for data in datas:
                data_te = data.get("_source")
                # print(time_now - datetime.fromisoformat(data_te.get("timestamp")).timestamp())
                if (time_now - datetime.datetime.fromisoformat(data_te.get("timestamp")).timestamp()) <= (60 * 1):
                    distributions.append({
                        "from": data_te.get("from"),
                        "to": data_te.get("to"),
                        "timestamp": data_te.get("timestamp")
                    })

            resData = {
                "devices": devices,
                "distributions": distributions
            }

            return Response(resData, status=status.HTTP_200_OK)
        except Exception as e:
            return Response(str(e), status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class DistribtionGraphView(APIView):
    def get(self, request):
        # try:
        elastic_server = os.environ.get(
            "ELASTIC_SERVER", "http://192.168.10.162:7772")
        # elastic_server = "http://192.168.10.162:7772"
        req = requests.get(
            f"{elastic_server}/demo-kc/_search?q=type_log:distribution&size=500&sort=timestamp:desc", )
        if req.status_code!=200:
            datas=[]
        else:
            datas = req.json().get("hits").get("hits")
        # print("99999999999999999", datas)
        devices = Devices.objects.all()
        distributions = []
        time_now = datetime.datetime.now(timezone.utc).timestamp()
        pairs = []

        edge_te = []
        node_te = [
            {
                "id": "centerserver",
                "label": "Center Server",
                "title": "Center Server",
                "shape": "circularImage",
                "image": f"http://192.168.10.162:7774/media/centerserver.jpg",
                "size": 40,
                "color": {
                    "border": "black",
                    "hover": {
                        "border": "black",
                    },
                    "hightlight": {
                        "border": "black",
                    },
                },
                "borderWidth": 2,
                "font": {
                    "color": "white",
                },
            },
        ]

        for data in datas:
            data_te = data.get("_source")
            # print(time_now - datetime.fromisoformat(data_te.get("timestamp")).timestamp())
            if (time_now - datetime.datetime.fromisoformat(data_te.get("timestamp")).timestamp()) <= (30):
                if [data_te.get("from"), data_te.get("to")] not in pairs:
                    pairs.append([data_te.get("from"), data_te.get("to")])
                    from_analyzer = f"analyzer_{'_'.join(str(data_te.get('from')).split('.'))}"
                    to_analyzer = f"analyzer_{'_'.join(str(data_te.get('to')).split('.'))}"
                    edge_te.append(
                        {"from": from_analyzer, "to": to_analyzer, "arrows": "to", "color": "red", "value": 1})
                    distributions.append(data_te.get("from"))
        print("------------", distributions)
        for device in devices:
            id_te = f"analyzer_{'_'.join(str(device.ip).split('.'))}"
            print("___________________________________", id_te)
            if device.ip in distributions:
                node_te.append({
                    "id": id_te,
                    "label": f"{device.ip}",
                    "title": f"{device.ip}",
                    "shape": "circularImage",
                    "image": f"http://192.168.10.162:7774/media/device.png",
                    "size": 20,
                    "color": {
                        "border": "red",
                        "hover": {
                            "border": "red",
                        },
                        "hightlight": {
                            "border": "red",
                        },
                    },
                    "borderWidth": 5,
                    "font": {
                        "color": "red",
                    },
                })
            else:
                node_te.append({
                    "id": id_te,
                    "label": f"{device.ip}",
                    "title": f"{device.ip}",
                    "shape": "circularImage",
                    "image": f"http://1/media/device.png",
                    "size": 20,
                    "color": {
                        "border": "black",
                        "hover": {
                            "border": "black",
                        },
                        "hightlight": {
                            "border": "black",
                        },
                    },
                    "borderWidth": 2,
                    "font": {
                        "color": "white",
                    },
                })
            edge_te.append({"from": id_te, "to": "centerserver"})

        resData = {
            "nodes": node_te,
            "edges": edge_te
        }

        return Response(resData, status=status.HTTP_200_OK)
        # except Exception as e:
        #     return Response(str(e), status=status.HTTP_500_INTERNAL_SERVER_ERROR)
class BlacklistClearView(APIView):
    def get(self, request):
        try:
            BlackListIP.objects.all().delete()
            return Response({
                "message": "Clear blacklist successfull"
            }, status=status.HTTP_200_OK)
        except Exception as err:
            return Response({
                "message": str(err)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class WhitelistClearView(APIView):
    def get(self, request):
        try:
            WhiteListIP.objects.all().delete()
            return Response({
                "message": "Clear whitelist successfull"
            }, status=status.HTTP_200_OK)
        except Exception as err:
            return Response({
                "message": str(err)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class BlacklistAllView(APIView):
    def get(self, request):
        try:
            blacklists = list(BlackListIP.objects.values())
            return Response({
                'count': len(blacklists),
                'results': blacklists
            }, status=status.HTTP_200_OK)
        except Exception as e:
            return Response(str(e), status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class WhitelistAllView(APIView):
    def get(self, request):
        try:
            whitelists = list(WhiteListIP.objects.values())
            return Response({
                'count': len(whitelists),
                'results': whitelists
            }, status=status.HTTP_200_OK)
        except Exception as e:
            return Response(str(e), status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(("POST", ))
def _send_telegram(request, *args, **kwargs):
    try:
        data = request.data
        if data.get("device_ip", None) == None:
            time_auto = int(get_env("AUTO_SEND_TELE_SECOND", 60 * 60 * 24))
            if time_auto == -1:
                time_auto = 60 * 60 * 24
            # print("time111111111111111111111111111111111, ", time_auto)
            days = time_auto // (60 * 60 * 24)
            time_auto -= (60 * 60 * 24 * days)
            hours = time_auto // (60 * 60)
            time_auto -= (60 * 60 * hours)
            mins = time_auto // (60)
            time_auto -= (60 * mins)
            secs = time_auto
            analyzers = Devices.objects.all()
            time_start = datetime.datetime.strftime(datetime.datetime.now() - timedelta(days=days, hours=hours, minutes=mins, seconds=secs), '%Y-%m-%d %H:%M:%S.%fZ')
            time_end = datetime.datetime.strftime(datetime.datetime.now(), '%Y-%m-%d %H:%M:%S.%fZ')
            print("===============================================")
            print(time_start, time_end)
            print("===============================================")
            for analyzer in analyzers:
                # result = Alerts.objects.all()
                result = Alerts.objects.all().filter(timestamp__range=(time_start, time_end))
                            
                body = f'''Trong vòng {days} ngày {hours} giờ {mins} phút {secs} giây vừa qua, đã có tổng cộng {len(result)} cuộc tấn công vào mạng lưới thiết bị IoT của bạn. Vui lòng kiểm tra email để biết thêm chi tiết.'''
                recipients = analyzer.telegram_id
                # if len(result) > 0:
                print("+++++++++++++++++++++++++++++++++++++++++++++")
                tele.send_to_telegram(message=body, chat_id=recipients)
        else:
            device_ip = data.get("device_ip")
            device = Devices.objects.filter(ip=device_ip).first()
            body = data.get("body")
            recipients = device.telegram_id
            tele.send_to_telegram(message=body, chat_id=recipients)
        return Response({
            'message': "Sent telegram"
        }, status=status.HTTP_200_OK)
    except Exception as e:
        logging.error(e)
        return Response(data={'status': False, 'Log bug': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)